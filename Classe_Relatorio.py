from docx import Document
from docx.shared import Inches
import pandas as pd
import matplotlib.pyplot as plt
import os
import matplotlib
from datetime import datetime
matplotlib.use('Agg')

class Relatorio:
    def __init__(self, dataset):
        self.dataset = dataset
        self.df_filtrado = dataset._df.copy()

    def adicionar_grafico(self, figura, doc, largura=5):
        plt.tight_layout()
        plt.savefig(figura)
        plt.close()
        doc.add_picture(figura, width=Inches(largura))
        os.remove(figura)

    def grafico_top5_bar(self, df, coluna, titulo, nome_arquivo, doc):
        df.groupby("Country")[coluna].mean().sort_values(ascending=False).head(5).plot(kind="bar")
        plt.title(titulo)
        plt.ylabel(coluna)
        self.adicionar_grafico(nome_arquivo, doc)

    def grafico_media_ano(self, df, coluna, titulo, nome_arquivo, doc, tipo='line'):
        df.groupby("Year")[coluna].mean().plot(kind=tipo, title=titulo)
        plt.ylabel(coluna)
        self.adicionar_grafico(nome_arquivo, doc)

    def grafico_pizza(self, df, coluna, titulo, nome_arquivo, doc):
        df[coluna].value_counts().plot.pie(autopct='%1.1f%%', startangle=90)
        plt.title(titulo)
        plt.ylabel("")
        self.adicionar_grafico(nome_arquivo, doc)

    def histograma(self, df, coluna, titulo, nome_arquivo, doc):
        df[coluna].plot.hist(bins=20, color='green')
        plt.title(titulo)
        plt.xlabel(coluna)
        self.adicionar_grafico(nome_arquivo, doc)

    def gerar_word(self, nome_arquivo="Relatorio.docx"):
        doc = Document()
        
        doc.add_paragraph("Aluno: Jeremias António Dinzinga")
        doc.add_paragraph("Curso: Ciência de dados e Inteligência Artificial - IPG")
        doc.add_paragraph("Aluno nº: 1710393")
        doc.add_paragraph("Projeto: Relatório de Expectativa de Vida")
        doc.add_paragraph(f"Data de Geração: {datetime.today().strftime('%d/%m/%Y')}")
        doc.add_paragraph("")
        
        doc.add_heading('Relatório Completo de Expectativa de Vida', 0)

        df = self.df_filtrado

        doc.add_heading("1. Top 5 Países com Maior e Menor Expectativa de Vida", level=1)
        for ano in [2013, 2015]:
            doc.add_heading(f"Ano: {ano}", level=2)
            dados_ano = df[df["Year"] == ano]
            if dados_ano.empty:
                doc.add_paragraph("Sem dados para este ano.")
                continue

            maiores = dados_ano[['Country', 'Life expectancy ']].sort_values(by='Life expectancy ', ascending=False).head(5)
            menores = dados_ano[['Country', 'Life expectancy ']].sort_values(by='Life expectancy ').head(5)

            # Tabela: Maior expectativa de vida
            doc.add_paragraph("Top 5 - Maior expectativa de vida:")
            tabela_maiores = doc.add_table(rows=1, cols=2)
            tabela_maiores.style = 'Table Grid'
            hdr_cells = tabela_maiores.rows[0].cells
            hdr_cells[0].text = 'País'
            hdr_cells[1].text = 'Expectativa de Vida'
            for _, row in maiores.iterrows():
                row_cells = tabela_maiores.add_row().cells
                row_cells[0].text = str(row['Country'])
                row_cells[1].text = f"{row['Life expectancy ']:.2f}"

            doc.add_paragraph("")  # espaço entre as tabelas

            # Tabela: Menor expectativa de vida
            doc.add_paragraph("Top 5 - Menor expectativa de vida:")
            tabela_menores = doc.add_table(rows=1, cols=2)
            tabela_menores.style = 'Table Grid'
            hdr_cells = tabela_menores.rows[0].cells
            hdr_cells[0].text = 'País'
            hdr_cells[1].text = 'Expectativa de Vida'
            for _, row in menores.iterrows():
                row_cells = tabela_menores.add_row().cells
                row_cells[0].text = str(row['Country'])
                row_cells[1].text = f"{row['Life expectancy ']:.2f}"

            doc.add_paragraph("")  # espaço final entre seções

        doc.add_heading("2. Top 5 Expectativa de Vida - 2015", level=1)
        top5_2015 = df[df["Year"] == 2015].nlargest(5, "Life expectancy ")
        top5_2015.plot(x="Country", y="Life expectancy ", kind="bar")
        plt.title("Top 5 Expectativa de Vida - 2015")
        plt.ylabel("Expectativa de Vida")
        self.adicionar_grafico("graf2.png", doc)

        doc.add_heading("3. Expectativa de Vida Média ao Longo dos Anos", level=1)
        self.grafico_media_ano(df, "Life expectancy ", "Expectativa de Vida Média por Ano", "graf1.png", doc)

        doc.add_heading("4. Distribuição por Estado de Desenvolvimento", level=1)
        self.grafico_pizza(df, "Status", "Distribuição por Status", "graf3.png", doc)

        doc.add_heading("5. Distribuição da Expectativa de Vida", level=1)
        self.histograma(df, "Life expectancy ", "Histograma da Expectativa de Vida", "graf4.png", doc)

        doc.add_heading("6. País com Maior Gasto em Saúde", level=1)
        gasto = df.groupby('Country')["Total expenditure"].mean().sort_values(ascending=False).head(1)
        doc.add_paragraph(gasto.to_string())

        doc.add_heading("7. Top 5 Países com Maior Gasto em Saúde", level=1)
        self.grafico_top5_bar(df, "Total expenditure", "Top 5 Gasto em Saúde", "graf5.png", doc)

        doc.add_heading("8. País com Maior Vacinação Infantil", level=1)
        vac = df.groupby('Country')["percentage expenditure"].mean().sort_values(ascending=False).head(1)
        doc.add_paragraph(vac.to_string())

        doc.add_heading("9. Top 5 Países com Maior Vacinação Infantil", level=1)
        self.grafico_top5_bar(df, "Diphtheria ", "Top 5 Vacinação Infantil", "graf6.png", doc)

        doc.add_heading("10. Evolução da Expectativa de Vida", level=1)
        self.grafico_media_ano(df, "Life expectancy ", "Evolução da Expectativa de Vida", "grafico_linha.png", doc)

        doc.add_heading("11. Gasto em Saúde Médio por Ano", level=1)
        self.grafico_media_ano(df, "Total expenditure", "Gasto em Saúde Médio por Ano", "grafico_gasto.png", doc, tipo="bar")

        doc.add_heading("12. Gráfico de Pizza: Status dos Países", level=1)
        self.grafico_pizza(df, "Status", "Distribuição por Status", "grafico_pizza.png", doc)

        doc.add_heading("13. Conclusão", level=1)
        doc.add_paragraph(
            "Este relatório apresenta um panorama completo sobre os dados de expectativa de vida, "
            "com destaque para os países com melhores indicadores, variações anuais e estatísticas relevantes."
        )

        doc.save(nome_arquivo)
        print(f"Relatório Word gerado: {nome_arquivo}")

    def converter_para_pdf(self, doc_path="relatorio.docx", pdf_path="relatorio.pdf"):
        from docx2pdf import convert
        convert(doc_path, pdf_path)
        print(f"Relatório PDF gerado: {pdf_path}")
